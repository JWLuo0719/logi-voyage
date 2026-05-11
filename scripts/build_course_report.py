from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
COURSE_DIR = ROOT / "课程信息"
OUTPUT_PATH = COURSE_DIR / "个人博客网站课程设计报告.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold


def format_paragraph(paragraph, line_spacing=1.5, space_before=0, space_after=6):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)


def add_text(paragraph, text, size=12, bold=False, east_asia="宋体", latin="Times New Roman"):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)
    return run


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = Inches(0.28)
    format_paragraph(paragraph, line_spacing=1.5, space_after=6)
    add_text(paragraph, text, size=12)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {level}"]
    format_paragraph(paragraph, line_spacing=1.2, space_before=12, space_after=6)
    add_text(paragraph, text, size=16 if level == 1 else 14, bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    format_paragraph(paragraph, line_spacing=1.35, space_after=4)
    add_text(paragraph, "• ", size=12)
    add_text(paragraph, text, size=12)
    return paragraph


def add_code_block(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(line)
        set_run_font(run, east_asia="等线", latin="Courier New", size=10)


def set_table_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, line_spacing=1.25, space_after=2)
    add_text(paragraph, text, size=11, bold=bold)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def set_column_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def ensure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
      style = doc.styles[style_name]
      style.font.name = "Times New Roman"
      style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
      style.font.size = Pt(size)
      style.font.bold = True

    if "Code Block" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(10)


def build_cover(doc):
    today = datetime.now()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(36)
    add_text(title, "湘潭大学", size=28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    subtitle.paragraph_format.space_after = Pt(60)
    add_text(subtitle, "《Web前端开发》课程设计报告", size=22, bold=True)

    fields = [
        ("题    目", "个人博客网站开发"),
        ("学    院", "计算机学院"),
        ("专    业", "待填写"),
        ("学    号", "待填写"),
        ("姓    名", "待填写"),
        ("任课教师", "杨喜喜"),
        ("完成时间", f"{today.year} 年 {today.month} 月"),
    ]
    for index, (label, value) in enumerate(fields):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(1.7)
        paragraph.paragraph_format.space_before = Pt(8 if index < 2 else 4)
        paragraph.paragraph_format.space_after = Pt(8)
        add_text(paragraph, f"{label}：", size=14, bold=True)
        add_text(paragraph, value, size=14)

    doc.add_page_break()


def build_toc(doc):
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.paragraph_format.space_after = Pt(18)
    add_text(toc, "目  录", east_asia="黑体", size=16, bold=True)

    toc_items = [
        "一、课程设计目的",
        "二、课程设计背景与主要技术工具介绍",
        "三、课程设计详情",
        "3.1 项目概述与目录结构",
        "3.2 课程要求完成情况",
        "3.3 主要页面设计与功能说明",
        "3.4 响应式设计与统一风格",
        "3.5 核心代码说明",
        "3.6 稳定性与运行效果分析",
        "四、遇到的问题及解决方法",
        "五、心得体会与收获",
    ]

    for item in toc_items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.45)
        format_paragraph(paragraph, line_spacing=1.35, space_after=4)
        add_text(paragraph, item, size=12)

    doc.add_page_break()


def build_body_content(doc):
    add_heading(doc, "一、课程设计目的", level=1)
    add_body_paragraph(doc, "本次课程设计以“个人博客网站开发”为题，目标是综合运用 HTML、CSS 和 JavaScript 完成一个多页面前端项目。通过课程设计，我希望把课堂上学习的页面结构语义化、响应式布局、DOM 交互、样式组织和多页面跳转等知识串联起来，形成一个能够独立展示、独立维护并且具备一定审美表达的完整网站。")
    add_body_paragraph(doc, "除了完成基础页面搭建，本项目还强调统一的视觉风格、较好的内容组织方式以及较清晰的用户浏览路径，使网站不仅“能运行”，也具备实际展示和持续扩展的价值。")

    add_heading(doc, "二、课程设计背景与主要技术工具介绍", level=1)
    add_body_paragraph(doc, "个人博客网站是前端课程设计中较适合综合练习的题目，因为它既包含首页、列表页、详情页和个人资料页等典型页面，也包含文章筛选、动态渲染、参数跳转和响应式适配等常见前端场景。相比只做单页静态页面，博客网站更能体现页面之间的数据复用、结构规划和样式统一能力。")
    add_body_paragraph(doc, "本项目采用纯前端技术实现，不依赖后端服务，主要技术工具如下：")
    add_bullet(doc, "HTML：负责四个核心页面的结构搭建，使用 header、main、section、article、nav 等语义化标签组织内容。")
    add_bullet(doc, "CSS：负责全站视觉风格、卡片布局、按钮样式、色彩系统以及响应式断点适配。")
    add_bullet(doc, "JavaScript：负责文章数据管理、首页推荐文章渲染、列表页搜索与分类筛选、详情页根据 URL 参数切换内容，以及移动端导航交互。")
    add_bullet(doc, "Visual Studio Code：作为主要开发工具，用于页面编写、样式调整和脚本调试。")
    add_bullet(doc, "Chrome DevTools：用于调试响应式布局、检查元素样式和验证页面交互。")
    add_bullet(doc, "Git / GitHub：用于版本管理和项目代码备份。")

    add_heading(doc, "三、课程设计详情", level=1)

    add_heading(doc, "3.1 项目概述与目录结构", level=2)
    add_body_paragraph(doc, "本项目名称为 Logi Voyage，是一个个人主页型博客网站。网站围绕“前端学习记录、AI 工具实战、个人内容入口”三个方向展开，采用统一的浅色玻璃拟态卡片风格，并通过公共数据文件驱动首页、文章列表页和文章详情页的内容展示。")
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    set_column_widths(table, [Inches(1.5), Inches(1.25), Inches(3.75)])
    headers = ["模块/文件", "作用", "说明"]
    for idx, title in enumerate(headers):
        set_table_cell(table.rows[0].cells[idx], title, bold=True)

    rows = [
        ("index.html", "首页", "展示网站简介、账号入口、推荐文章和站点概览"),
        ("posts.html", "文章列表页", "提供文章搜索、分类筛选和列表浏览"),
        ("post.html", "文章详情页", "根据 URL 参数动态展示对应文章内容"),
        ("profile.html", "个人资料页", "展示个人简介、技能方向、账号信息和成长时间线"),
        ("styles/site.css", "全局样式", "统一颜色变量、布局体系、卡片组件和响应式规则"),
        ("styles/home.css", "页面扩展样式", "补充首页、资料页等特殊布局样式"),
        ("scripts/data.js", "数据源", "集中存放文章标题、摘要、正文和封面等数据"),
        ("scripts/site.js", "全局脚本", "导航切换、日期格式化、文章卡片生成"),
        ("scripts/posts.js", "列表交互", "实现搜索和分类筛选"),
        ("scripts/post.js", "详情渲染", "实现文章详情页动态切换"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_table_cell(cells[idx], value)

    add_heading(doc, "3.2 课程要求完成情况", level=2)
    add_body_paragraph(doc, "对照老师给出的“个人博客网站开发”题目要求，当前项目已经完成课程要求中的硬性部分，并具备提交课程设计系统源码的基础。完成情况如下表所示。")
    checklist = doc.add_table(rows=1, cols=3)
    set_table_borders(checklist)
    set_column_widths(checklist, [Inches(2.3), Inches(1.1), Inches(3.1)])
    for idx, title in enumerate(["课程要求", "完成情况", "对应实现"]):
        set_table_cell(checklist.rows[0].cells[idx], title, bold=True)
    checklist_rows = [
        ("使用 HTML、CSS 和 JavaScript 实现博客网站基本结构", "已完成", "四个页面均由 HTML 搭建，样式由 CSS 统一管理，交互与渲染由 JavaScript 完成"),
        ("至少包含首页、文章列表页、文章详情页和个人资料页", "已完成", "项目包含 index.html、posts.html、post.html、profile.html 四个核心页面"),
        ("实现响应式设计，确保不同设备显示良好", "已完成", "site.css 与 home.css 中设置了 960px 和 760px 等断点，移动端导航采用折叠菜单"),
        ("实现多页面、统一风格、网页元素丰富、界面优美", "已完成", "全站使用统一色彩变量、卡片组件、按钮、排版和内容模块，包含卡片、时间线、筛选按钮等元素"),
        ("源码书写规范，必要地方有注释", "已完成", "脚本按模块拆分，关键函数已补充说明性注释，便于后续维护"),
    ]
    for row in checklist_rows:
        cells = checklist.add_row().cells
        for idx, value in enumerate(row):
            set_table_cell(cells[idx], value)

    add_heading(doc, "3.3 主要页面设计与功能说明", level=2)
    add_body_paragraph(doc, "（1）首页：首页用于建立访客的第一印象，包含网站标题、个人简介、头像、账号入口、站点概览和推荐文章等模块。首页通过读取统一的文章数据，只展示最新三篇文章，保证信息量适中且维护方便。")
    add_body_paragraph(doc, "（2）文章列表页：文章列表页承担网站内容总入口的作用，支持按关键词搜索和按分类筛选，用户可以快速定位到自己感兴趣的文章主题。该页面在不刷新页面的情况下即可完成列表更新，交互逻辑清晰。")
    add_body_paragraph(doc, "（3）文章详情页：详情页通过读取地址栏中的 id 参数，动态匹配 scripts/data.js 中的对应文章对象，再渲染标题、摘要、正文、推荐阅读和侧边信息。该实现方式避免为每篇文章单独制作一个 HTML 文件，提高了扩展性。")
    add_body_paragraph(doc, "（4）个人资料页：资料页主要用于展示作者身份、技能方向、内容平台入口和成长时间线，与首页形成补充关系。页面模块包括头像介绍、账号卡片、技能标签和时间线等，增强了网站的完整度。")

    add_heading(doc, "3.4 响应式设计与统一风格", level=2)
    add_body_paragraph(doc, "为了保证网站在电脑端和移动端都有较好的浏览体验，项目采用了基于 Flex 和 Grid 的响应式布局方案。全局样式文件通过媒体查询在 960px 以下将多列布局压缩为单列布局，在 760px 以下启用折叠导航，并进一步收紧外边距、标题字号和卡片内边距。")
    add_body_paragraph(doc, "在视觉风格上，项目通过 CSS 变量统一定义背景色、强调色、阴影、圆角和文本颜色，再在首页、列表页、详情页和资料页中复用同一套卡片组件、按钮样式和段落节奏，保证了网站风格统一。网页元素方面，网站包含导航栏、页眉横幅、卡片组、标签、按钮、搜索框、分类筛选、推荐阅读、时间线和账号入口等模块，整体界面较为完整。")

    add_heading(doc, "3.5 核心代码说明", level=2)
    add_body_paragraph(doc, "项目的核心实现思路是“统一数据源 + 多页面复用组件 + 根据页面职责分别渲染”。下面选取几段代表性代码进行说明。")
    add_body_paragraph(doc, "1）全局脚本中统一生成文章卡片，避免首页和列表页重复书写同类 HTML：")
    add_code_block(doc, [
        'window.createPostCard = function createPostCard(post) {',
        '  return [',
        '    \'<article class="post-card">\',',
        '    \'  <h3><a href="post.html?id=\' + post.id + \'">\' + post.title + "</a></h3>\',',
        '    "  <p>" + post.summary + "</p>",',
        '  ].join("");',
        '};',
    ])
    add_body_paragraph(doc, "2）文章列表页通过关键词和分类共同控制筛选结果，提高文章浏览效率：")
    add_code_block(doc, [
        'function render() {',
        '  var keyword = searchInput.value.trim().toLowerCase();',
        '  var filtered = posts.filter(function (post) {',
        '    var matchesCategory = activeCategory === "全部" || post.category === activeCategory;',
        '    var haystack = [post.title, post.summary, post.category].join(" ").toLowerCase();',
        '    return matchesCategory && (!keyword || haystack.indexOf(keyword) !== -1);',
        '  });',
        '}',
    ])
    add_body_paragraph(doc, "3）文章详情页通过 URL 参数匹配文章 id，实现“一页模板 + 多篇文章”的动态切换：")
    add_code_block(doc, [
        'var params = new URLSearchParams(window.location.search);',
        'var currentId = params.get("id") || (posts[0] && posts[0].id);',
        'var currentPost = posts.find(function (post) { return post.id === currentId; });',
        'if (!currentPost) {',
        '  container.innerHTML = "...没有找到这篇文章...";',
        '  return;',
        '}',
    ])
    add_body_paragraph(doc, "4）移动端导航切换通过 class 和 aria-expanded 同步控制，提高了小屏设备下的可用性：")
    add_code_block(doc, [
        'toggle.addEventListener("click", function () {',
        '  var open = nav.classList.toggle("is-open");',
        '  toggle.setAttribute("aria-expanded", open ? "true" : "false");',
        '});',
    ])
    add_body_paragraph(doc, "从整体实现上看，这些代码既满足了课程设计的基本交互需求，也体现了模块拆分、数据复用和页面解耦的思路。")

    add_heading(doc, "3.6 稳定性与运行效果分析", level=2)
    add_body_paragraph(doc, "由于本项目为纯静态前端网站，不涉及数据库读写、接口请求或复杂计算，因此整体运行稳定性较高。网站的主要交互逻辑集中在页面渲染、筛选和跳转，脚本复杂度适中，执行效率能够满足课程设计要求。")
    add_body_paragraph(doc, "在运行效果方面，项目能够完成多页面跳转、文章动态切换、首页推荐展示、分类筛选和移动端导航折叠等功能。页面布局在桌面端和移动端均能保持较清晰的视觉层级，符合课程设计“布局风格统一、网页元素丰富、界面优美”的要求。需要说明的是，目前项目尚未接入后端管理功能，文章内容维护仍依赖前端数据文件手动更新，但这不影响课程设计题目要求的达成。")

    add_heading(doc, "四、遇到的问题及解决方法", level=1)
    add_body_paragraph(doc, "问题一：多个页面都需要展示文章信息，如果每个页面分别手写卡片和正文，后期维护成本较高。解决方法是将文章标题、分类、摘要、正文等信息集中放入 scripts/data.js，再由不同页面脚本按需读取和渲染，从而减少重复代码。")
    add_body_paragraph(doc, "问题二：文章详情页数量会随着内容增加而增长，如果每篇文章都单独新建一个详情 HTML 页面，结构会很臃肿。解决方法是保留一个通用的 post.html 模板页面，通过 URL 参数匹配不同文章 id，实现动态切换。")
    add_body_paragraph(doc, "问题三：桌面端多列布局在手机上容易拥挤，导航链接也不便点击。解决方法是在 CSS 中设置断点，将多列布局调整为单列布局，同时增加移动端导航按钮，通过 JavaScript 控制展开和收起。")
    add_body_paragraph(doc, "问题四：不同页面在配色、按钮、圆角和阴影处理上如果缺少统一规则，会导致网站整体观感松散。解决方法是将背景、强调色、圆角和阴影提取成 CSS 变量，并在全站卡片和按钮组件中复用。")

    add_heading(doc, "五、心得体会与收获", level=1)
    add_body_paragraph(doc, "通过这次课程设计，我对“完整前端项目”和“单个练习页面”之间的区别有了更清晰的认识。单个页面更强调局部实现，而多页面网站更强调信息结构、页面关系、风格统一和后续维护方式。")
    add_body_paragraph(doc, "在实现过程中，我进一步巩固了 HTML 语义化结构、CSS Grid/Flex 布局、媒体查询、DOM 操作和 URL 参数读取等知识，并逐步形成了先拆结构、再做样式、最后补交互的开发习惯。")
    add_body_paragraph(doc, "此外，我也认识到前端项目不仅要关注“功能有没有做出来”，还要关注代码是否易于维护、页面是否有一致的视觉规则、用户浏览路径是否清晰。后续如果继续扩展本项目，我会考虑增加文章发布管理、更多专题页面以及更完善的作品展示能力。")


def add_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = ""
        add_text(footer, "《Web前端开发》课程设计报告", size=10)


def main():
    doc = Document()
    ensure_styles(doc)
    build_cover(doc)
    build_toc(doc)
    build_body_content(doc)
    add_footer(doc)
    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
